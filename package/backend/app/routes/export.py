from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.models import OptimizationSegment, OptimizationSession
from app.routes.optimization import get_current_user

router = APIRouter(prefix="/export", tags=["export"])


@router.get("/sessions/{session_id}/docx")
async def export_docx(
    session_id: str,
    card_key: str,
    db: Session = Depends(get_db),
):
    """导出会话结果为 Word 文档"""
    user = get_current_user(card_key, db)

    session = db.query(OptimizationSession).filter(
        OptimizationSession.session_id == session_id,
        OptimizationSession.user_id == user.id,
    ).first()

    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    if session.status != "completed":
        raise HTTPException(status_code=400, detail="会话未完成，无法导出")

    segments = (
        db.query(OptimizationSegment)
        .filter(OptimizationSegment.session_id == session.id)
        .order_by(OptimizationSegment.segment_index)
        .all()
    )

    doc = Document()

    title = doc.add_heading("AI 论文润色结果", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"会话 ID：{session.session_id}")
    doc.add_paragraph(f"创建时间：{session.created_at.strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    for seg in segments:
        final_text = (
            seg.user_edited_text
            or seg.enhanced_text
            or seg.polished_text
            or seg.original_text
        )

        heading_text = f"段落 {seg.segment_index + 1}"
        if seg.is_title:
            heading_text += "（标题）"
        doc.add_heading(heading_text, level=2)

        doc.add_heading("原文", level=3)
        original_paragraph = doc.add_paragraph(seg.original_text or "")
        for run in original_paragraph.runs:
            run.font.size = Pt(11)

        doc.add_heading("润色结果", level=3)
        final_paragraph = doc.add_paragraph(final_text or "")
        for run in final_paragraph.runs:
            run.font.size = Pt(11)

        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"optimized_{session_id}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
